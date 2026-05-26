"""Word template filling — placeholder text replacement + results table writing."""

from docx.document import Document as _DocType
from docx.oxml.ns import qn as _qn  # re-export shim for clarity
from docx.oxml.ns import qn
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


def _iter_block_items(parent):
    if isinstance(parent, _DocType):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        return
    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def _replace_in_paragraph(paragraph: Paragraph, mapping: dict[str, str]) -> None:
    if not paragraph.runs:
        return
    full_text = "".join(r.text or "" for r in paragraph.runs)
    if not any(k in full_text for k in mapping):
        return
    new_text = full_text
    for old, new in mapping.items():
        new_text = new_text.replace(old, new)
    paragraph.runs[0].text = new_text
    for r in paragraph.runs[1:]:
        r.text = ""


def _replace_in_blocks(parent, mapping: dict[str, str]) -> None:
    for block in _iter_block_items(parent):
        if isinstance(block, Paragraph):
            _replace_in_paragraph(block, mapping)
        elif isinstance(block, Table):
            for row in block.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        _replace_in_paragraph(p, mapping)


def replace_text_everywhere(doc, mapping: dict[str, str]) -> None:
    """Replace placeholder text in body, headers, and footers of `doc`."""
    _replace_in_blocks(doc, mapping)
    for section in doc.sections:
        for p in section.header.paragraphs:
            _replace_in_paragraph(p, mapping)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        _replace_in_paragraph(p, mapping)
        for p in section.footer.paragraphs:
            _replace_in_paragraph(p, mapping)
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        _replace_in_paragraph(p, mapping)


# Header tokens that identify the standard (reference) and actual (calibration) columns.
# Japanese: 基準値 / 実際の値. English fallbacks: Standard / Actual / Reference / Measured.
_STD_TOKENS = ("基準値", "Standard", "Reference", "Ref")
_ACT_TOKENS = ("実際の値", "実測", "Actual", "Measured", "Cal")


def _matches_any(text: str, tokens: tuple[str, ...]) -> bool:
    lower = text.lower()
    for t in tokens:
        if t in text or t.lower() in lower:
            return True
    return False


def find_results_table(doc) -> tuple[Table, int, int, list[int]]:
    """Locate the calibration results table.

    Returns (table, col_std, col_act, data_row_indices). The data rows are
    positionally aligned with the ordered targets [-40, 5, 40] — caller maps
    target → row by zipping.
    """
    for table in doc.tables:
        if not table.rows:
            continue
        header = [c.text.strip() for c in table.rows[0].cells]
        col_std = next((i for i, h in enumerate(header) if _matches_any(h, _STD_TOKENS)), -1)
        col_act = next((i for i, h in enumerate(header) if _matches_any(h, _ACT_TOKENS)), -1)
        if col_std < 0 or col_act < 0:
            continue
        data_rows = list(range(1, len(table.rows)))
        return table, col_std, col_act, data_rows
    raise RuntimeError("Could not find calibration results table in template")


def _set_cell(cell, text: str) -> None:
    if cell.paragraphs:
        p = cell.paragraphs[0]
        if p.runs:
            p.runs[0].text = text
            for r in p.runs[1:]:
                r.text = ""
            return
    cell.text = text


def fill_results_table(
    doc, ordered_values: list[tuple[float | None, float | None]]
) -> None:
    """Fill the calibration results table.

    `ordered_values` is a list of (standard, actual) tuples, one per data row,
    in the same order as the data rows of the table (i.e. -40, 5, 40 if using
    the default matcher.TARGETS).
    """
    table, col_std, col_act, data_rows = find_results_table(doc)
    for i, row_idx in enumerate(data_rows):
        if i >= len(ordered_values):
            break
        std, act = ordered_values[i]
        row = table.rows[row_idx]
        _set_cell(row.cells[col_std], "" if std is None else f"{std:.1f}")
        _set_cell(row.cells[col_act], "" if act is None else f"{act:.1f}")
