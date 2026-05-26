from docx import Document

from ite_api.calibration.docx_filler import (
    fill_results_table,
    find_results_table,
    replace_text_everywhere,
)


def test_replace_text_everywhere_updates_cert_number(template_docx, tmp_path):
    doc = Document(str(template_docx))
    replace_text_everywhere(doc, {"0000001700": "0000001999"})
    out = tmp_path / "out.docx"
    doc.save(str(out))
    saved = Document(str(out))
    body_text = "\n".join(p.text for p in saved.paragraphs)
    assert "0000001999" in body_text
    assert "0000001700" not in body_text


def test_replace_text_everywhere_handles_serial_in_table(template_docx, tmp_path):
    doc = Document(str(template_docx))
    replace_text_everywhere(doc, {"190124110002417": "190124110099999"})
    out = tmp_path / "out.docx"
    doc.save(str(out))
    saved = Document(str(out))
    # Serial lives in table 0
    table_text = " ".join(
        c.text for t in saved.tables for r in t.rows for c in r.cells
    )
    assert "190124110099999" in table_text
    assert "190124110002417" not in table_text


def test_replace_text_everywhere_handles_dates(template_docx, tmp_path):
    doc = Document(str(template_docx))
    replace_text_everywhere(doc, {
        "2026年3月4日": "2026年4月14日",
        "2026年3月6日": "2026年4月15日",
    })
    out = tmp_path / "out.docx"
    doc.save(str(out))
    saved = Document(str(out))
    body_text = "\n".join(p.text for p in saved.paragraphs)
    assert "2026年4月14日" in body_text
    assert "2026年4月15日" in body_text
    assert "2026年3月4日" not in body_text
    assert "2026年3月6日" not in body_text


def test_find_results_table_locates_the_setpoint_table(template_docx):
    doc = Document(str(template_docx))
    table, col_std, col_act, data_rows = find_results_table(doc)
    assert table is not None
    assert col_std != col_act
    assert len(data_rows) == 3


def test_fill_results_table_writes_each_row(template_docx, tmp_path):
    doc = Document(str(template_docx))
    fill_results_table(doc, [
        (-40.10, -40.05),
        (5.02, 5.04),
        (39.95, 40.01),
    ])
    out = tmp_path / "out.docx"
    doc.save(str(out))
    saved = Document(str(out))
    table, col_std, col_act, data_rows = find_results_table(saved)
    for i, (std, act) in enumerate(["-40.10", "5.02", "39.95"]
                                   and [("-40.10", "-40.05"), ("5.02", "5.04"), ("39.95", "40.01")]):
        row = table.rows[data_rows[i]]
        assert std in row.cells[col_std].text
        assert act in row.cells[col_act].text
