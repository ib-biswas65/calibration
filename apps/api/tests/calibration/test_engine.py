from datetime import datetime
from pathlib import Path

from docx import Document

from ite_api.calibration.cal_loader import load_workbook
from ite_api.calibration.engine import (
    BatchConfig,
    RunConfig,
    SetpointWindow,
    run_calibration,
    run_one_logger,
)
from ite_api.calibration.ref_loader import load_ref_auto

_WIDE_START = datetime(1900, 1, 1)
_WIDE_END = datetime(2999, 12, 31, 23, 59)
_SETPOINTS = [
    SetpointWindow(target=-40.0, start=_WIDE_START, end=_WIDE_END),
    SetpointWindow(target=5.0,   start=_WIDE_START, end=_WIDE_END),
    SetpointWindow(target=40.0,  start=_WIDE_START, end=_WIDE_END),
]


def test_run_one_logger_produces_docx(workbook_xlsx, reference_csv, template_docx, tmp_path):
    cfg = RunConfig(
        cert_no="0000001999",
        serial="190124110099999",
        test_date_jp="2026年4月14日",
        doc_date_jp="2026年4月15日",
        template_path=template_docx,
        output_dir=tmp_path,
        setpoints=_SETPOINTS,
    )
    wb, names = load_workbook(workbook_xlsx)
    ref_df = load_ref_auto(reference_csv)
    out_path = run_one_logger(cfg, sheet_name=names[0], wb=wb, ref_df=ref_df)
    assert out_path.exists()
    assert out_path.suffix == ".docx"
    saved = Document(str(out_path))
    body_text = "\n".join(p.text for p in saved.paragraphs)
    assert "0000001999" in body_text
    table_text = " ".join(
        c.text for t in saved.tables for r in t.rows for c in r.cells
    )
    assert "190124110099999" in table_text


def test_run_calibration_writes_one_docx_per_sheet(
    workbook_xlsx, reference_csv, template_docx, tmp_path
):
    cfg = BatchConfig(
        start_cert_no="0000002000",
        cert_width=10,
        test_date_jp="2026年4月14日",
        doc_date_jp="2026年4月15日",
        template_path=template_docx,
        calibration_xlsxs=[workbook_xlsx],
        reference_csvs=[reference_csv],
        output_dir=tmp_path / "out",
        setpoints=_SETPOINTS,
    )
    written = run_calibration(cfg)
    assert len(written) > 0
    for p in written:
        assert p.exists() and p.suffix == ".docx"
    nums = sorted(int(p.name.split("_")[2]) for p in written)
    assert nums == list(range(nums[0], nums[0] + len(nums)))
