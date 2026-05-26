"""Word template filling logic — preserves all docx manipulation from calibration.py."""

import os
import re
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


def iter_block_items(parent):
    """Iterate through paragraphs and tables in a document or cell."""
    if isinstance(parent, DocxDocument):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        parent_elm = parent._element
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def replace_text_in_paragraph(paragraph, mapping: dict):
    """Replace text across all runs in a paragraph.

    Word often splits a single visible string across multiple XML <w:r> runs
    due to internal formatting, spell-check marks, or edit history.
    For example, '0000001644' might be stored as runs ['00000016', '44'].

    This function concatenates all run texts, performs the replacement on the
    combined string, and then redistributes the result back to the original
    runs so that all existing formatting (font, size, bold, color) is preserved.
    """
    runs = paragraph.runs
    if not runs:
        return

    # Step 1: Build the full paragraph text from all runs
    full_text = "".join(r.text for r in runs)

    # Check if any replacement is needed at all (fast path)
    needs_replace = False
    for old in mapping:
        if old in full_text:
            needs_replace = True
            break
    if not needs_replace:
        return

    # Step 2: Perform all replacements on the combined text
    for old, new in mapping.items():
        full_text = full_text.replace(old, new)

    # Step 3: Redistribute the new text back across the original runs.
    # The first run gets all the text; remaining runs are emptied.
    # This preserves the first run's formatting for the entire replaced text.
    runs[0].text = full_text
    for run in runs[1:]:
        run.text = ""


def _replace_in_blocks(parent, mapping: dict):
    """Replace placeholder text in all paragraphs and table cells of a parent element."""
    for block in iter_block_items(parent):
        if isinstance(block, Paragraph):
            replace_text_in_paragraph(block, mapping)
        elif isinstance(block, Table):
            for row in block.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_text_in_paragraph(p, mapping)


def replace_text_everywhere(doc, mapping: dict):
    """Replace placeholder text in the entire document: body, headers, and footers.

    Japanese calibration certificate templates commonly place the cert number
    and serial in headers/footers. This function ensures those are replaced too.
    """
    # 1. Document body
    _replace_in_blocks(doc, mapping)

    # 2. Section headers and footers
    seen = set()
    for section in doc.sections:
        for hf in (section.header, section.footer,
                   section.first_page_header, section.first_page_footer,
                   section.even_page_header, section.even_page_footer):
            if hf is None or hf.is_linked_to_previous:
                continue
            hf_id = id(hf._element)
            if hf_id in seen:
                continue
            seen.add(hf_id)
            for p in hf.paragraphs:
                replace_text_in_paragraph(p, mapping)
            for tbl in hf.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            replace_text_in_paragraph(p, mapping)


def set_cell_text_centered(cell, text: str):
    """Set cell text while preserving formatting, centered."""
    if cell.paragraphs and cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].text = text
        for run in cell.paragraphs[0].runs[1:]:
            run.text = ""
    else:
        cell.text = text

    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tcPr = cell._tc.get_or_add_tcPr()
        vAlign = OxmlElement("w:vAlign")
        vAlign.set(qn("w:val"), "center")
        tcPr.append(vAlign)


def find_calibration_table(doc):
    """Find the calibration results table and return (table, col_std, col_act, row_map)."""
    targets = ["-40.00", "5.00", "40.00"]
    for tbl in doc.tables:
        header_row = None
        for ri, row in enumerate(tbl.rows):
            cleaned = ["".join(c.text.split()) for c in row.cells]
            if any("基準値" in t for t in cleaned) and any(
                "実際の値" in t for t in cleaned
            ):
                header_row = ri
                break
        if header_row is None:
            continue

        headers = [re.sub(r"\s+", "", c.text) for c in tbl.rows[header_row].cells]
        try:
            col_std = next(i for i, h in enumerate(headers) if "基準値" in h)
            col_act = next(i for i, h in enumerate(headers) if "実際の値" in h)
        except StopIteration:
            continue

        row_map = {}
        for ri in range(header_row + 1, len(tbl.rows)):
            txt = tbl.rows[ri].cells[col_std].text.strip()
            m = re.search(r"-?\d+\.?\d*", txt)
            if m:
                val = float(m.group(0))
                # Match to nearest target (within ±1°C tolerance)
                for t in targets:
                    if abs(val - float(t)) <= 1.0:
                        row_map[t] = ri
                        break

        if row_map:
            return tbl, col_std, col_act, row_map

    return None, None, None, None


def fill_certificate(
    template_path: Path,
    output_path: Path,
    cert_no: str,
    serial: str,
    values: dict,
    targets: list,
    template_serial: str,
    test_date_jp: str,
    doc_date_jp: str,
    template_testdate: str,
    template_docdate: str,
    template_cert_no: str = "0000001644",
):
    """
    Create a certificate from the template, filling in all values.

    Args:
        template_path: Path to .docx template
        output_path: Path to save the generated certificate
        cert_no: The certificate number string (zero-padded)
        serial: Logger serial number
        values: dict mapping target_temp → (ref_value, cal_value)
        targets: list of target temperatures [-40.0, 5.0, 40.0]
        template_serial: Placeholder serial in template
        test_date_jp: New test date (Japanese format)
        doc_date_jp: New document date (Japanese format)
        template_testdate: Placeholder test date in template
        template_docdate: Placeholder doc date in template
        template_cert_no: The cert number string currently in the template
                          (will be replaced with cert_no). Configurable so
                          it works with any template, not just one with "0000001644".
    """
    doc = Document(str(template_path))

    # Serial number in product table (Table 0, Row 1, Cell 2 = 型番)
    product_tbl = doc.tables[0]
    serial_cell = product_tbl.rows[1].cells[2]
    if serial_cell.paragraphs and serial_cell.paragraphs[0].runs:
        serial_cell.paragraphs[0].runs[0].text = serial
        for run in serial_cell.paragraphs[0].runs[1:]:
            run.text = ""
    else:
        serial_cell.text = serial

    # Build replacement mapping — cert number key is now fully configurable
    mapping = {
        template_cert_no: cert_no,
        template_serial: serial,
        template_testdate: test_date_jp,
        template_docdate: doc_date_jp,
    }
    replace_text_everywhere(doc, mapping)

    # Fill calibration table
    tbl, col_std, col_act, row_map = find_calibration_table(doc)

    summary_entries = []
    if tbl is not None:
        for target in targets:
            key = f"{target:.2f}"
            if key in row_map:
                ri = row_map[key]
                ref_val, cal_val = values[target]
                set_cell_text_centered(tbl.rows[ri].cells[col_std], f"{ref_val:.2f}")
                set_cell_text_centered(tbl.rows[ri].cells[col_act], f"{cal_val:.2f}")

                summary_entries.append(
                    {
                        "Certificate No": cert_no,
                        "Serial": serial,
                        "Target (°C)": target,
                        "Reference (°C)": ref_val,
                        "Actual (°C)": cal_val,
                        "Difference (°C)": round(cal_val - ref_val, 2),
                        "Within ±0.5": (
                            "Yes" if abs(cal_val - ref_val) <= 0.5 else "No"
                        ),
                    }
                )

    # Save the certificate
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    return summary_entries
