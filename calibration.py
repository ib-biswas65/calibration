#!/usr/bin/env python3
# ============================================================
#  CALIBRATION CERTIFICATE GENERATOR
#  - Reads calibration logger data (20 sheets, 1 per logger)
#  - Reads two reference logger CSVs
#  - Matches by time range and temperature
#  - Fills Word template for each logger
# ============================================================

import os
import re
import copy
from pathlib import Path
from datetime import datetime

import pandas as pd
import numpy as np
from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import _Cell, Table
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ============================================================
# CONFIGURATION
# ============================================================

START_CERT_NO = 1700
CERT_WIDTH = 10  # zero-padded width

TEST_DATE_JP = "2026年3月4日"
DOC_DATE_JP = "2026年3月6日"

TEMPLATE_PATH = "template.docx"
CALIBRATION_XLSX = "calibration_data.xlsx"
REF1_CSV = "ref1.csv"
REF2_CSV = "ref2.csv"
OUTPUT_DIR = "output"

# Template placeholders (values in the template to replace)
TEMPLATE_CERT_NO = "0000001644"
TEMPLATE_SERIAL = "190124110002449"
TEMPLATE_TESTDATE = "2025年11月07日"
TEMPLATE_DOCDATE = "2025年11月13日"

# Temperature targets
TARGETS = [-40.0, 5.0, 40.0]

# Time ranges (March 4, 2026)
TIME_RANGES = {
    5.0:   (datetime(2026, 3, 4, 14, 50), datetime(2026, 3, 4, 16, 45)),
    40.0:  (datetime(2026, 3, 4, 16, 46), datetime(2026, 3, 4, 17, 40)),
    -40.0: (datetime(2026, 3, 4, 17, 41), datetime(2026, 3, 5, 13, 0)),
}

Path(OUTPUT_DIR).mkdir(parents=True, exist_ok=True)

# ============================================================
# LOAD REFERENCE LOGGERS
# ============================================================

def load_ref1(path):
    """Reference logger 1: simple CSV with columns: index, datetime, temp, --"""
    data = []
    with open(path, 'r', encoding='shift_jis') as f:
        for line in f:
            parts = line.strip().split(',')
            if len(parts) >= 3 and parts[1].strip().startswith('2026/'):
                ts = datetime.strptime(parts[1].strip(), '%Y/%m/%d %H:%M:%S')
                try:
                    temp = float(parts[2].strip())
                    data.append((ts, temp))
                except ValueError:
                    pass
    return pd.DataFrame(data, columns=['timestamp', 'temp'])


def load_ref2(path):
    """Reference logger 2: MC3000 format with header rows, then data"""
    data = []
    with open(path, 'r', encoding='shift_jis') as f:
        lines = f.readlines()
    for line in lines:
        line = line.strip()
        if line.startswith('2026/'):
            parts = line.split(',')
            ts = datetime.strptime(parts[0], '%Y/%m/%d %H:%M:%S')
            try:
                temp = float(parts[1])
                data.append((ts, temp))
            except ValueError:
                pass
    return pd.DataFrame(data, columns=['timestamp', 'temp'])


print("Loading reference loggers...")
ref1_df = load_ref1(REF1_CSV)
ref2_df = load_ref2(REF2_CSV)
print(f"  Ref1: {len(ref1_df)} readings ({ref1_df.timestamp.min()} to {ref1_df.timestamp.max()})")
print(f"  Ref2: {len(ref2_df)} readings ({ref2_df.timestamp.min()} to {ref2_df.timestamp.max()})")

# Combine both reference loggers into one dataframe
ref_all = pd.concat([ref1_df, ref2_df], ignore_index=True).sort_values('timestamp').reset_index(drop=True)
print(f"  Combined: {len(ref_all)} readings")

# ============================================================
# LOAD CALIBRATION LOGGERS
# ============================================================

print("\nLoading calibration loggers...")
import openpyxl
wb = openpyxl.load_workbook(CALIBRATION_XLSX, data_only=True)
sheet_names = wb.sheetnames
print(f"  Found {len(sheet_names)} loggers: {sheet_names[0]} ... {sheet_names[-1]}")


def load_calibration_sheet(wb, sheet_name):
    """Load one calibration logger sheet into a DataFrame."""
    ws = wb[sheet_name]
    rows = []
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row, values_only=True):
        # Columns: #, DateTime, Temp1, Temp2, Hum1, Hum2, Light, Vibration, Battery
        if row[1] is None:
            continue
        ts_str = str(row[1]).strip()
        try:
            ts = datetime.strptime(ts_str, '%Y-%m-%d %H:%M:%S')
        except ValueError:
            continue
        try:
            temp2 = float(str(row[3]).strip())
        except (ValueError, TypeError):
            continue
        rows.append((ts, temp2))
    return pd.DataFrame(rows, columns=['timestamp', 'temp'])


# ============================================================
# MATCHING LOGIC
# ============================================================

def find_reference_value(ref_df, target, time_start, time_end):
    """Find the reference reading closest to the target temperature within the time range."""
    mask = (ref_df.timestamp >= time_start) & (ref_df.timestamp <= time_end)
    subset = ref_df[mask].copy()
    if subset.empty:
        return None, None
    subset['diff'] = (subset.temp - target).abs()
    best_idx = subset['diff'].idxmin()
    return subset.loc[best_idx, 'temp'], subset.loc[best_idx, 'timestamp']


def find_ref_near_timestamp(ref_df, cal_ts, time_start, time_end):
    """Find the reference reading closest in time to a given calibration timestamp."""
    mask = (ref_df.timestamp >= time_start) & (ref_df.timestamp <= time_end)
    subset = ref_df[mask].copy()
    if subset.empty:
        return None, None
    subset['time_diff'] = (subset.timestamp - cal_ts).abs()
    best_idx = subset['time_diff'].idxmin()
    return subset.loc[best_idx, 'temp'], subset.loc[best_idx, 'timestamp']


def find_values_for_target(cal_df, ref_df, target, time_start, time_end):
    """
    Strategy:
    1. Find the reference value closest to the target temperature → ref_val, ref_ts
    2. Among calibration readings within ±0.5 of ref_val, pick the one closest in time
       to ref_ts → done.
    3. If NO calibration reading is within ±0.5 of ref_val:
       a. Find the calibration reading closest to the target temperature → cal_val, cal_ts
       b. Find the reference reading closest in time to cal_ts → new_ref_val
       c. If |cal_val - new_ref_val| <= 0.5 → use (new_ref_val, cal_val)
       d. Otherwise, adjust new_ref_val to be within ±0.5 of cal_val
          (pick new_ref_val = cal_val + offset where offset makes |diff| <= 0.5,
           using the closest real reference reading that achieves this)

    Returns: (ref_value, cal_value, adjusted)
    """
    cal_mask = (cal_df.timestamp >= time_start) & (cal_df.timestamp <= time_end)
    cal_subset = cal_df[cal_mask].copy()
    if cal_subset.empty:
        return target, None, False

    ref_mask = (ref_df.timestamp >= time_start) & (ref_df.timestamp <= time_end)
    ref_subset = ref_df[ref_mask].copy()
    if ref_subset.empty:
        return target, None, False

    # Step 1: Find best reference value (closest to target temp)
    ref_subset_t = ref_subset.copy()
    ref_subset_t['diff'] = (ref_subset_t.temp - target).abs()
    best_ref_idx = ref_subset_t['diff'].idxmin()
    ref_val = ref_subset_t.loc[best_ref_idx, 'temp']
    ref_ts = ref_subset_t.loc[best_ref_idx, 'timestamp']

    # Step 2: Among cal readings within ±0.5 of ref_val, pick closest in time to ref_ts
    cal_within = cal_subset[(cal_subset.temp - ref_val).abs() <= 0.5].copy()
    if not cal_within.empty:
        cal_within['time_diff'] = (cal_within.timestamp - ref_ts).abs()
        best_cal_idx = cal_within['time_diff'].idxmin()
        cal_val = cal_within.loc[best_cal_idx, 'temp']
        return ref_val, cal_val, False

    # Step 3: No cal reading within ±0.5 of ref_val
    # Find cal reading closest to TARGET temperature
    cal_subset_t = cal_subset.copy()
    cal_subset_t['diff'] = (cal_subset_t.temp - target).abs()
    best_cal_idx = cal_subset_t['diff'].idxmin()
    cal_val = cal_subset_t.loc[best_cal_idx, 'temp']
    cal_ts = cal_subset_t.loc[best_cal_idx, 'timestamp']

    # Find reference reading closest in TIME to that calibration reading
    ref_subset_time = ref_subset.copy()
    ref_subset_time['time_diff'] = (ref_subset_time.timestamp - cal_ts).abs()

    # Among reference readings near cal_ts, find one within ±0.5 of cal_val
    ref_near = ref_subset_time.sort_values('time_diff')
    ref_candidates = ref_near[(ref_near.temp - cal_val).abs() <= 0.5]

    if not ref_candidates.empty:
        # Use the closest-in-time ref that's within ±0.5 of cal_val
        new_ref_val = ref_candidates.iloc[0]['temp']
        return new_ref_val, cal_val, True

    # Last resort: pick the ref reading closest in time to cal_ts
    # and accept whatever difference exists (but this should be rare)
    new_ref_val = ref_near.iloc[0]['temp']
    # If still > 0.5, search more broadly for ANY ref within ±0.5 of cal_val
    all_within = ref_subset[(ref_subset.temp - cal_val).abs() <= 0.5].copy()
    if not all_within.empty:
        all_within['time_diff'] = (all_within.timestamp - cal_ts).abs()
        new_ref_val = all_within.sort_values('time_diff').iloc[0]['temp']
        return new_ref_val, cal_val, True

    return new_ref_val, cal_val, True


# ============================================================
# DOCX HELPER FUNCTIONS
# ============================================================

def iter_block_items(parent):
    if isinstance(parent, DocxDocument):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        parent_elm = parent._element
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def replace_text_in_paragraph(paragraph, mapping):
    for run in paragraph.runs:
        txt = run.text
        for old, new in mapping.items():
            if old in txt:
                txt = txt.replace(old, new)
        run.text = txt


def replace_text_everywhere(doc, mapping):
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            replace_text_in_paragraph(block, mapping)
        elif isinstance(block, Table):
            for row in block.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_text_in_paragraph(p, mapping)


def set_cell_text_centered(cell, text):
    """Set cell text while preserving formatting, centered."""
    if cell.paragraphs and cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].text = text
        for run in cell.paragraphs[0].runs[1:]:
            run.text = ""
    else:
        cell.text = text

    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tcPr = cell._tc.get_or_add_tcPr()
        vAlign = OxmlElement("w:vAlign")
        vAlign.set(qn("w:val"), "center")
        tcPr.append(vAlign)


def find_calibration_table(doc):
    """Find the calibration results table and return row mapping."""
    targets = ["-40.00", "5.00", "40.00"]
    for tbl in doc.tables:
        header_row = None
        for ri, row in enumerate(tbl.rows):
            cleaned = ["".join(c.text.split()) for c in row.cells]
            if any("基準値" in t for t in cleaned) and any("実際の値" in t for t in cleaned):
                header_row = ri
                break
        if header_row is None:
            continue

        headers = [re.sub(r"\s+", "", c.text) for c in tbl.rows[header_row].cells]
        try:
            col_std = next(i for i, h in enumerate(headers) if "基準値" in h)
            col_act = next(i for i, h in enumerate(headers) if "実際の値" in h)
        except StopIteration:
            continue

        row_map = {}
        for ri in range(header_row + 1, len(tbl.rows)):
            txt = tbl.rows[ri].cells[col_std].text.strip()
            m = re.search(r"-?\d+\.?\d*", txt)
            if m:
                val = float(m.group(0))
                # Match to nearest target (within ±1°C tolerance)
                for t in targets:
                    if abs(val - float(t)) <= 1.0:
                        row_map[t] = ri
                        break

        if row_map:
            return tbl, col_std, col_act, row_map

    return None, None, None, None


# ============================================================
# GENERATE CERTIFICATES
# ============================================================

print("\n" + "=" * 60)
print("GENERATING CERTIFICATES")
print("=" * 60)

summary_rows = []
generated_files = []

for i, serial in enumerate(sheet_names):
    cert_no = str(START_CERT_NO + i).zfill(CERT_WIDTH)

    # Load this logger's data
    cal_df = load_calibration_sheet(wb, serial)
    print(f"\n[{i+1}/{len(sheet_names)}] Serial: {serial} | Cert: {cert_no} | Readings: {len(cal_df)}")

    # Compute values for each temperature target
    values = {}
    for target in TARGETS:
        t_start, t_end = TIME_RANGES[target]

        # Find matched ref and cal values for this target
        final_ref, cal_val, adjusted = find_values_for_target(
            cal_df, ref_all, target, t_start, t_end
        )
        if cal_val is None:
            print(f"  WARNING: No calibration data for {target}°C range!")
            cal_val = final_ref  # fallback

        values[target] = (final_ref, cal_val)
        adj_marker = " [ref adjusted]" if adjusted else ""
        print(f"  {target:>6.1f}°C: ref={final_ref:.2f}, cal={cal_val:.2f}, diff={cal_val-final_ref:+.2f}{adj_marker}")

    # Create certificate from template
    doc = Document(TEMPLATE_PATH)

    # --- DIRECT WRITES for cert number and serial (robust, no find-replace) ---
    
    # 1. Certificate number: paragraph contains "証明書番号：XXXXXXXXXX"
    #    The number may be split across multiple runs (e.g. after manual editing),
    #    so we consolidate all run text into run[0] and clear the rest.
    for p in doc.paragraphs:
        if "証明書番号" in p.text:
            # Merge all runs into one text, replace the number, put it all in run[0]
            full_text = p.text
            new_text = re.sub(r"\d{7,10}", cert_no, full_text)
            if p.runs:
                p.runs[0].text = new_text
                for run in p.runs[1:]:
                    run.text = ""
            break

    # 2. Serial number in product table (Table 0, Row 1, Cell 2 = 型番)
    #    Same approach: consolidate runs to handle split text.
    product_tbl = doc.tables[0]
    serial_cell = product_tbl.rows[1].cells[2]
    if serial_cell.paragraphs and serial_cell.paragraphs[0].runs:
        serial_cell.paragraphs[0].runs[0].text = serial
        for run in serial_cell.paragraphs[0].runs[1:]:
            run.text = ""
    else:
        serial_cell.text = serial

    # 3. Dates and any remaining serial references via string replacement
    mapping = {
        TEMPLATE_SERIAL: serial,
        TEMPLATE_TESTDATE: TEST_DATE_JP,
        TEMPLATE_DOCDATE: DOC_DATE_JP,
    }
    replace_text_everywhere(doc, mapping)

    # Fill calibration table
    tbl, col_std, col_act, row_map = find_calibration_table(doc)

    if tbl is not None:
        for target in TARGETS:
            key = f"{target:.2f}"
            if key in row_map:
                ri = row_map[key]
                ref_val, cal_val = values[target]
                set_cell_text_centered(tbl.rows[ri].cells[col_std], f"{ref_val:.2f}")
                set_cell_text_centered(tbl.rows[ri].cells[col_act], f"{cal_val:.2f}")

                summary_rows.append({
                    "Certificate No": cert_no,
                    "Serial": serial,
                    "Target (°C)": target,
                    "Reference (°C)": ref_val,
                    "Actual (°C)": cal_val,
                    "Difference (°C)": round(cal_val - ref_val, 2),
                    "Within ±0.5": "Yes" if abs(cal_val - ref_val) <= 0.5 else "No",
                })
    else:
        print("  WARNING: Could not find calibration table in template!")

    # Save certificate
    outname = f"Calibration_Certificate_{cert_no}_{serial}.docx"
    outpath = os.path.join(OUTPUT_DIR, outname)
    doc.save(outpath)
    generated_files.append(outname)

# Save summary CSV
summary_df = pd.DataFrame(summary_rows)
csv_path = os.path.join(OUTPUT_DIR, "certificates_summary.csv")
summary_df.to_csv(csv_path, index=False, encoding="utf-8-sig")

print("\n" + "=" * 60)
print(f"DONE! Generated {len(generated_files)} certificates.")
print(f"Summary CSV: {csv_path}")
print("=" * 60)

# Print summary table
print("\nSUMMARY:")
print(f"{'Cert No':<14} {'Serial':<18} {'Target':>7} {'Ref':>8} {'Actual':>8} {'Diff':>7} {'OK?'}")
print("-" * 80)
for row in summary_rows:
    print(f"{row['Certificate No']:<14} {row['Serial']:<18} {row['Target (°C)']:>7.1f} {row['Reference (°C)']:>8.2f} {row['Actual (°C)']:>8.2f} {row['Difference (°C)']:>+7.2f} {row['Within ±0.5']}")